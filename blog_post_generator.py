from dotenv import load_dotenv
import os
import json
from openai import OpenAI
from docx import Document

class BlogPostGenerator:
    def __init__(self, title):
        self.title = title
        self.document = Document()
        self.document.add_heading(title, level=1)
        self.load_environment_variables()
        self.client = OpenAI(api_key=os.getenv("API_KEY"))

    def load_environment_variables(self):
        load_dotenv()

    def save_document(self):
        self.document.save(f'{self.title}.docx')

    def save_text_file(self, content, mode='a'):
        with open(f'{self.title}.txt', mode, encoding='utf-8') as file:
            file.write(content + '\n\n')  # Add extra newline for readability

    def generate_outline(self):
        blog_input = self._format_blog_input()
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "assistant", "content": blog_input}]
        )
        outline = response.choices[0].message.content
        print("Outline done")
        self.document.add_paragraph(outline)
        self.save_document()
        
        print(outline)
        self.parse_outline_json(outline)

    def _format_blog_input(self):
        return f'''
        I run a coding blog. I am writing a blog post titled "{self.title}". Can you give me a general outline of the post in json format in this manner: 

        {{
        "blog_outline":
        [
            {{
            "name_of_section": "____",
            "subsections": ["___,____,____,....."]
            }},
            ...
        ]
        }}
        '''

    def parse_outline_json(self, outline):
        data = json.loads(outline)
        for each in data["blog_outline"]:
            subsections_str = ', '.join(each["subsections"])
            query = f"{each['name_of_section']}: {subsections_str}"
            self.query_openai_for_section(query, outline)
            print(f"Querying: {query}")
        print("Done. Please check")

    def query_openai_for_section(self, query, outline):
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "assistant", "content": f"I am writing a blog post titled {self.title} with this outline {outline}, expand in a detailed informative manner providing code snippets where necessary about this specific subsection: {query}. Please give it in a formatted layout with h2 headings starting with '#'; h3 headings starting with '##' and code snippets starting with ```. "}]
        )
        sub_output = response.choices[0].message.content
        self.document = Document(f'{self.title}.docx')
        self.document.add_heading(query, level=2)
        self.document.add_paragraph(sub_output)
        self.save_document()
        self.save_text_file(query + '\n' + sub_output)  # Save each section to the text file

def text_to_html(self, input_path=None, output_path=None):
    if input_path is None:
        input_path = f'{self.title}.txt'
    if output_path is None:
        output_path = f'{self.title}.html'

    with open(input_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    html_lines = ['<html>', '<head><title>Document</title></head>', '<body>']
    inside_code_block = False

    for line in lines:
        if inside_code_block:
            # Check if the line signifies the end of a code block
            if line.strip() == "```":
                inside_code_block = False
                html_lines.append('</code></pre>')  # Close the code block
                continue  # Skip adding ``` to output
            else:
                # Replace leading spaces with &nbsp; for indentation within code blocks
                leading_spaces = len(line) - len(line.lstrip(' '))
                space_replacement = '&nbsp;' * leading_spaces
                html_lines.append(space_replacement + line.lstrip())
        else:
            # Check if the line signifies the start of a code block, ignoring ```python
            if line.strip() in ("```", "```python"):
                inside_code_block = True
                html_lines.append('<pre><code>')  # Open a new code block
                continue  # Skip adding ``` or ```python to output
            else:
                stripped_line = line.strip()
                if stripped_line.startswith("#"):
                    level = stripped_line.count("#") + 1
                    heading = stripped_line.lstrip("#").strip()
                    html_lines.append(f'<h{level}>{heading}</h{level}>')
                elif stripped_line:
                    html_lines.append(f'<p>{stripped_line}</p>')

    html_lines.append('</body></html>')

    with open(output_path, 'w', encoding='utf-8') as file:
        file.write('\n'.join(html_lines))


# Usage
blog_title = "Introduction to OOPS in Python"
blog_post_generator = BlogPostGenerator(blog_title)
blog_post_generator.generate_outline()
# Once the outline and content are generated and saved to a text file
blog_post_generator.text_to_html()  # Converts the text file to an HTML file
