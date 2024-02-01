from dotenv import load_dotenv
import os
import json
from openai import OpenAI
from docx import Document

# Load environment variables
load_dotenv()
API_KEY=os.getenv("API_KEY")
# Setting API client
client = OpenAI(api_key=API_KEY)

# Defining the blog title
blog_title = "File IO in Python"

# Making a word docx file and saving it with a title and a name which is the blog_title

document = Document()
document.add_heading(blog_title, level=1)
document.save(f'{blog_title}.docx')

# Get input for blog
blog_input = '''
I run a coding blog. I am writing a blog post titled "{}". Can you give me a general outline of the post in json format in this manner: 

{{
  "blog_outline":
  [
    {{
      "name_of_section": "____",
      "subsections": ["___,____,____,....."]
    }},
    ...
  ]
}}
'''.format(blog_title)

def outline():
    response = client.chat.completions.create(
        model="gpt-4",
        messages= [{"role": "assistant", "content": blog_input},]
        )

    outline = (response.choices[0].message.content) 
    print("outline done")
    # saving outline to docx file
    # Open the existing document
    document = Document(f'{blog_title}.docx')
    # Add the sub_output
    document.add_paragraph(outline)
    # Save the document
    document.save(f'{blog_title}.docx')
    print(outline)
    json_parse(outline)

def json_parse(outline):
    data = json.loads(outline)

    for each in data["blog_outline"]:
        subsections_str = ', '.join(each["subsections"])  # Joining the subsections into a single string
        query = f"{each['name_of_section']}: {subsections_str}"
        openai_query(query, outline)
        print(f"querying: {query}" )
    print("Done. Please check")

def openai_query(query, outline):
    response = client.chat.completions.create(
        model="gpt-4",
        messages= [{"role": "assistant", "content": f"I am writing a blog post titled {blog_title} with this outline {outline}, expand in a detailed informative manner providing code snippets where necessary about this specific subsection: {query} "},]
        )
    sub_output = response.choices[0].message.content
    # saving output to docx file
    # Open the existing document
    document = Document(f'{blog_title}.docx')
    # Add the sub_output
    document.add_heading(query, level=1) # This saves the subsection and subheading
    document.add_paragraph(sub_output)
    # Save the document
    document.save(f'{blog_title}.docx')


outline()
